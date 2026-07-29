"""DOCX → raw JSON 변환. 단락/표/섹션 구조 추출."""
import json
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn


# 파일명 → 문서 유형(메타데이터) 힌트. 컬렉션 분기가 아니라 검색 필터·표시용 태그다.
# 매칭 실패 시 "general" — 어떤 공개 문서든 파이프라인에 넣을 수 있다.
CONTRACT_TYPE_MAP = {
    "공공계약 실무가이드": "general",
    "공공계약실무가이드": "general",
    "공공SW사업": "service",
    "건설엔지니어링": "construction",
}

# 섹션 제목으로 판단할 스타일 이름 또는 텍스트 패턴
HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "제목 1", "제목 2", "제목 3"}
HEADING_FONT_BOLD_MIN_SIZE = 14  # pt


def _is_heading(para) -> tuple[bool, int]:
    """단락이 섹션 제목인지 판단. (is_heading, level) 반환."""
    style_name = para.style.name if para.style else ""
    if style_name in HEADING_STYLES:
        level = int(style_name[-1]) if style_name[-1].isdigit() else 1
        return True, level

    # 스타일이 없는 경우 폰트 크기 + 굵기로 판단
    text = para.text.strip()
    if not text:
        return False, 0

    for run in para.runs:
        font = run.font
        size = font.size.pt if font.size else 0
        bold = font.bold
        if bold and size >= HEADING_FONT_BOLD_MIN_SIZE:
            return True, 2

    # □, ○, ◆ 등 특수기호로 시작하는 중제목 패턴
    if re.match(r"^[□○◆◇▶▷●■◉]", text) and len(text) < 60:
        return True, 3

    return False, 0


def _get_merged_cell_text(cell) -> str:
    """셀 내 모든 단락 텍스트를 합쳐 반환."""
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _parse_table(table, table_index: int) -> dict:
    """표를 dict로 변환. 병합셀은 첫 셀 값으로 대표."""
    rows_data = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            row_cells.append(_get_merged_cell_text(cell))
        rows_data.append(row_cells)

    if not rows_data:
        return {"table_index": table_index, "headers": [], "rows": []}

    # 첫 행을 헤더로 취급
    headers = rows_data[0]
    data_rows = []
    for row in rows_data[1:]:
        if len(row) == len(headers) and headers:
            row_dict = {headers[i]: row[i] for i in range(len(headers))}
        else:
            row_dict = {f"col_{i}": v for i, v in enumerate(row)}
        data_rows.append(row_dict)

    return {
        "table_index": table_index,
        "headers": headers,
        "row_count": len(data_rows),
        "col_count": len(headers),
        "rows": data_rows,
    }


def parse_docx(file_path: Path) -> dict:
    """DOCX 파일을 섹션/단락/표 구조의 dict로 파싱."""
    doc = Document(str(file_path))

    # 파일명에서 contract_type(메타 태그)·document_id 추론
    stem = file_path.stem
    contract_type = "general"
    for key, val in CONTRACT_TYPE_MAP.items():
        if key in stem:
            contract_type = val
            break
    year_m = re.search(r"(20\d\d)", stem)
    slug = re.sub(r"[^가-힣a-zA-Z0-9_]", "_", stem)[:30]
    document_id = f"{slug}_{year_m.group(1)}" if year_m and year_m.group(1) not in slug else slug

    sections = []
    current_section = None
    table_index = 0

    # 문서 본문 요소를 순서대로 순회 (단락 + 표 혼합)
    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    next_para = next(para_iter, None)
    next_table = next(table_iter, None)

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":  # 단락
            if next_para is None:
                continue
            para = next_para
            next_para = next(para_iter, None)

            text = para.text.strip()
            is_h, level = _is_heading(para)

            if is_h and text:
                section_id = f"{document_id}_s{len(sections):03d}"
                current_section = {
                    "section_id": section_id,
                    "title": text,
                    "level": level,
                    "paragraphs": [],
                    "tables": [],
                }
                sections.append(current_section)
            elif text:
                if current_section is None:
                    current_section = {
                        "section_id": f"{document_id}_s000",
                        "title": "서론",
                        "level": 1,
                        "paragraphs": [],
                        "tables": [],
                    }
                    sections.append(current_section)
                current_section["paragraphs"].append({
                    "style": para.style.name if para.style else "",
                    "text": text,
                })

        elif tag == "tbl":  # 표
            if next_table is None:
                continue
            table = next_table
            next_table = next(table_iter, None)

            parsed = _parse_table(table, table_index)
            table_index += 1

            if current_section is None:
                current_section = {
                    "section_id": f"{document_id}_s000",
                    "title": "서론",
                    "level": 1,
                    "paragraphs": [],
                    "tables": [],
                }
                sections.append(current_section)
            current_section["tables"].append(parsed)

    return {
        "document_id": document_id,
        "source_file": file_path.name,
        "contract_type": contract_type,
        "version": year_m.group(1) if year_m else "",
        "total_sections": len(sections),
        "total_tables": table_index,
        "sections": sections,
    }


if __name__ == "__main__":
    _root = Path(__file__).resolve().parents[2]
    src_dir = _root / "data" / "source_docs"
    out_dir = _root / "etl" / "data" / "raw"
    out_dir.mkdir(parents=True, exist_ok=True)

    if not src_dir.exists():
        raise SystemExit(f"소스 문서 디렉터리 없음: {src_dir} — 공개 문서 DOCX를 넣은 뒤 재실행")

    docx_files = sorted(src_dir.glob("*.docx"))
    if not docx_files:
        print(f"DOCX 없음: {src_dir}")
    for docx_file in docx_files:
        print(f"파싱 중: {docx_file.name}")
        result = parse_docx(docx_file)
        out_path = out_dir / f"raw_{result['document_id']}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"  → {out_path} (섹션: {result['total_sections']}, 표: {result['total_tables']})")
